#!/usr/bin/env python3
"""
Build a single docx revision sheet from all INTERVIEW_WALKTHROUGH.md files.

Structure:
  - Title page
  - Auto-paginating TOC (Word field — F9 to populate)
  - For each problem:
      * Heading 1 (picked up by TOC)
      * Tier badge
      * 1-page Quick Card (entities, patterns, signal, traps, soundbite)
      * Full walkthrough rendered from MD
  - Reference pages at the back:
      1. Pattern → Problem index
      2. Concurrency idiom → Problem index
      3. Company → Priority problems
      4. 45-minute round playbook
      5. Anti-pattern catalog
"""
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = Path(__file__).parent / "src/main/java/com/conceptcoding/interviewquestions/hello_all_questions"
OUTPUT = Path(__file__).parent / "LLD_Revision_Sheet.docx"

# Static index page numbers — populated from extract_page_numbers.py after
# rendering the docx via LibreOffice. When non-empty, the Contents page
# renders as a static two-column table (no F9 needed). When empty (or None),
# falls back to a Word TOC field placeholder.
#
# Update workflow:
#   1. set STATIC_INDEX_PAGES = {} and run this script
#   2. soffice --headless --convert-to pdf --outdir /tmp LLD_Revision_Sheet.docx
#   3. python3 extract_page_numbers.py  → copy the JSON output here
#   4. re-run this script — repeat steps 2-3 once more to check stability
STATIC_INDEX_PAGES = {
    "1. Parking Lot": 3,
    "2. Snake & Ladder": 21,
    "3. Vending Machine": 29,
    "4. LRU Cache": 44,
    "5. Connect Four": 58,
    "6. Splitwise": 76,
    "7. Movie Ticket Booking": 92,
    "8. Amazon Locker": 113,
    "9. Logging Service": 133,
    "10. Rate Limiter": 154,
    "11. Inventory Management": 175,
    "12. Notification System": 196,
    "13. Job Scheduler": 214,
    "14. Meeting Scheduler": 232,
    "15. Payment Gateway": 249,
    "16. URL Shortener": 266,
    "17. File System": 281,
    "18. Cab Booking": 303,
    "19. Chess": 313,
    "20. Elevator System": 332,
    "Pattern → Problem Index": 352,
    "Concurrency Idiom → Problem Index": 354,
    "Company → Priority Problems": 356,
    "45-Minute Round Playbook": 357,
    "Anti-Pattern Catalog": 359,
}

# ----------------------- Study sequence (curated) -----------------------
STUDY_SEQUENCE = [
    # ---- Foundation warm-ups ----
    ("parkinglot",       "Parking Lot",           "Tier 1 — Foundation"),
    ("snakeladder",      "Snake & Ladder",        "Tier 1 — Foundation"),
    ("vendingmachine",   "Vending Machine",       "Tier 1 — Foundation"),
    ("lrucache",         "LRU Cache",             "Tier 1 — Foundation"),
    ("connectfour",      "Connect Four",          "Tier 1 — Foundation"),
    # ---- Core company favorites ----
    ("splitwise",        "Splitwise",             "Tier 2 — Core"),
    ("movieticket",      "Movie Ticket Booking",  "Tier 2 — Core"),
    ("amazonlocker",     "Amazon Locker",         "Tier 2 — Core"),
    ("logger",           "Logging Service",       "Tier 2 — Core"),
    ("ratelimiter",      "Rate Limiter",          "Tier 2 — Core"),
    # ---- Advanced patterns + concurrency ----
    ("inventory",        "Inventory Management",  "Tier 3 — Advanced"),
    ("notification",     "Notification System",   "Tier 3 — Advanced"),
    ("jobscheduler",     "Job Scheduler",         "Tier 3 — Advanced"),
    ("meetingscheduler", "Meeting Scheduler",     "Tier 3 — Advanced"),
    ("paymentgateway",   "Payment Gateway",       "Tier 3 — Advanced"),
    # ---- Multi-pattern flagships ----
    ("urlshortener",     "URL Shortener",         "Tier 4 — Flagship"),
    ("filesystem",       "File System",           "Tier 4 — Flagship"),
    ("cabbooking",       "Cab Booking",           "Tier 4 — Flagship"),
    ("chess",            "Chess",                 "Tier 4 — Flagship"),
    ("elevator",         "Elevator System",       "Tier 4 — Flagship"),
]


# ----------------------- Quick Card content (per problem) -----------------------
# Each card is the distilled revision view — read this for 2nd-pass revision,
# flip to the full walkthrough only when you need detail.
QUICK_CARDS = {
    "parkinglot": {
        "entities": "ParkingLot, Floor, Spot (with VehicleSize), Vehicle, Ticket, ParkingLotService (facade)",
        "patterns_day1": "Facade (service), Tell-Don't-Ask (Spot.tryReserve), money in long cents",
        "patterns_step5": "Strategy (SpotLookupStrategy) — defended in Step 5, not pre-baked",
        "signal": "Per-spot synchronized for park-attempt CAS — NEVER a global lock on the lot. Lock granularity = throughput.",
        "traps": [
            "Global lock kills throughput under contention",
            "Storing money as double / float",
            "Caller branches on VehicleSize instead of letting Spot decide fit",
        ],
        "soundbite": "Spot has tryReserve → atomic per-spot CAS. Floor scans spots. Lot composes floors. The lock granularity defines whether this scales — per-spot synchronized stays correct without serializing the whole lot.",
    },
    "snakeladder": {
        "entities": "Player (id, position), Board (snakes, ladders, applyJump), Dice (Strategy), Game (turn loop)",
        "patterns_day1": "Strategy (Dice — for testability), Tell-Don't-Ask (Board.applyJump)",
        "patterns_step5": "Observer if 'stream moves to UI' asked; minimum-rolls-to-win as BFS on board-as-graph",
        "signal": "Board.applyJump encapsulates snake/ladder resolution; Game NEVER branches on snake vs. ladder. Tell, Don't Ask.",
        "traps": [
            "Putting snake/ladder if-else in Game instead of Board",
            "Flaky tests — hard-coding Math.random instead of Dice Strategy",
            "Forgetting overshoot rule (target > size → stay put)",
        ],
        "soundbite": "Game is a turn loop. Board is geometry — answers 'where do I end up after landing on N'. Dice Strategy on day 1 because I need a deterministic test-fake; without it every snake-bite scenario is flaky.",
    },
    "vendingmachine": {
        "entities": "VendingMachine (context), VendingMachineState (interface), NoCoinState, HasCoinState, DispensingState, Product, Inventory",
        "patterns_day1": "State (class-per-state — GoF), flyweight (3 state singletons), Strategy-ish for state behavior",
        "patterns_step5": "Strategy for refund / change-making policies; Observer for restock alerts",
        "signal": "Class-per-state because each state has DIFFERENT behavior for the same action. selectProduct on HasCoinState dispenses; on NoCoinState it rejects. Compare with enum-state-machines (PaymentGateway) where transitions are pure bookkeeping.",
        "traps": [
            "Reaching for enum-state-machine here — wrong fit because behavior differs per state",
            "Creating new state objects per transition (use 3 flyweights)",
            "Allowing mutation from invalid state (e.g., dispense() from NoCoinState)",
        ],
        "soundbite": "Each state IS a class because behavior differs per state — not just data. Context holds three flyweight state instances. selectProduct on HasCoinState transitions to DispensingState then chains to dispense(). This is the textbook GoF State pattern — earn it by showing distinct per-state behavior, not just labels.",
    },
    "lrucache": {
        "entities": "Cache<K,V> interface, LRUCache (HashMap + DLL), Node<K,V> (key, value, prev, next), sentinel head/tail",
        "patterns_day1": "Composition (HashMap + DLL), interface seam, sentinel nodes to kill null edges",
        "patterns_step5": "LinkedHashMap-accessOrder alternative (10-line variant); thread-safety via synchronized or ReentrantLock; LFU as separate strategy",
        "signal": "HashMap for O(1) key lookup + Doubly-Linked-List for O(1) eviction + sentinel head/tail = three O(1) operations (get, put, evict) with no null-checks in hot path.",
        "traps": [
            "Using LinkedList — O(n) move-to-front operation",
            "Forgetting to move-to-front on get (not just put)",
            "Not handling key UPDATES (must move existing node, not insert duplicate)",
        ],
        "soundbite": "HashMap gives O(1) key→node. DLL gives O(1) eviction at tail and promotion to head. Sentinel head/tail nodes eliminate every null-check in pointer manipulation. On every get OR put, the node moves to head. When size exceeds capacity, tail.prev evicts.",
    },
    "connectfour": {
        "entities": "Board (cols × rows), Cell (Empty/Red/Yellow), Player, Game (turn + status: IN_PROGRESS/WIN/DRAW)",
        "patterns_day1": "Enum for cell state, simple state machine (3 statuses), gravity logic on column drop",
        "patterns_step5": "Strategy for win-detection variants (4-in-row vs N-in-row); AI player via minimax",
        "signal": "Win-check ONLY from the last-dropped cell in 4 directions (horizontal / vertical / two diagonals). O(1) per move, not O(board). Cheap because the new piece is the only thing that can complete a line.",
        "traps": [
            "Scanning entire board after every move (O(board) instead of O(1))",
            "Forgetting gravity — piece floats mid-column",
            "Allowing drops on full columns",
        ],
        "soundbite": "Drop into column → gravity finds the lowest empty row. Win-check only from the new piece, 4 directions, count consecutive same-color. O(1) per move. Status transitions IN_PROGRESS → WIN (with winner) or DRAW (board full).",
    },
    "splitwise": {
        "entities": "User, Group, Expense, Split (SplitType: EQUAL/EXACT/PERCENTAGE), Balance (Map<owner, Map<owedTo, cents>>)",
        "patterns_day1": "Strategy (SplitStrategy per type), Tell-Don't-Ask, long cents, basis points for percentages",
        "patterns_step5": "Settle-up minimization (graph reduction — fewer net transfers); Observer on balance threshold",
        "signal": "Balance is a graph: Map<owner, Map<owedTo, cents>>. Cross-payments use ORDERED LOCK ACQUISITION (sort the two user IDs) to prevent deadlock. Remainder cents (equal-split of 10.01 among 3) go to the first payer — keeps sum exact.",
        "traps": [
            "doubles for money — silent rounding errors compound",
            "Locking two users in arbitrary order → deadlock on cross-payments",
            "Symmetric rounding of equal-split remainder (sum no longer matches expense)",
        ],
        "soundbite": "Expense produces Splits via Strategy. Balance graph updated atomically using ordered-key locking (smaller ID acquired first). Money is long cents. Percentages are basis points (10000 = 100%). Settle-up is a separate graph-reduction pass — not done eagerly on every expense.",
    },
    "movieticket": {
        "entities": "Theater, Show, Movie, Seat (status: AVAILABLE/HELD/BOOKED), Booking, User; Show owns Map<SeatId, SeatStatus>",
        "patterns_day1": "Per-seat synchronized for hold→book CAS, state machine for Seat, Strategy for pricing",
        "patterns_step5": "Hold timeout via ScheduledExecutorService, surge pricing, payment confirmation flow, fraud detection",
        "signal": "Two-phase booking: HOLD (with TTL) → CONFIRM. Per-seat synchronized makes hold atomic. Held seats auto-release via scheduled task. Two concurrent bookings on the same seat → exactly one wins the hold.",
        "traps": [
            "Locking the whole show (kills throughput)",
            "Not releasing HELD seats on payment-timeout",
            "Allowing direct AVAILABLE → BOOKED, skipping HELD",
        ],
        "soundbite": "Two-phase booking — hold-with-TTL, then confirm. Per-seat synchronized CAS gives atomic AVAILABLE → HELD. ScheduledExecutorService releases stale holds. The state machine enforces hold-before-book — direct AVAILABLE → BOOKED throws IllegalStateException.",
    },
    "amazonlocker": {
        "entities": "Locker (size: SMALL/MEDIUM/LARGE, status), LockerLocation, Package, ReservationCode, LockerService",
        "patterns_day1": "Per-locker synchronized tryReserve (atomic AVAILABLE→RESERVED), Strategy for size matching",
        "patterns_step5": "Observer for 'package picked up' notification, code expiry via ScheduledExecutor, audit log",
        "signal": "Smallest-locker-that-fits via Strategy. Per-locker tryReserve is the atomic AVAILABLE → RESERVED CAS — same idiom as Parking Lot's tryReserve. One reservation code maps 1:1 to one locker.",
        "traps": [
            "Reserving by 'pick first match' without atomic CAS",
            "Storing reservation codes plaintext — hash them",
            "Not expiring stale codes (locker held forever)",
        ],
        "soundbite": "Strategy picks the smallest-locker-that-fits. Per-locker synchronized tryReserve gives atomic AVAILABLE → RESERVED. Code expiry is a scheduled task that flips back to AVAILABLE. Pickup releases — handing the package back is the user's responsibility.",
    },
    "logger": {
        "entities": "Logger, LogLevel, LogMessage, LogAppender (Console/File/Remote — Strategy), LogFilter (Level/Keyword — Decorator)",
        "patterns_day1": "Strategy (appenders + filters), Decorator (filter wraps appender), async fan-out via BlockingQueue + worker thread",
        "patterns_step5": "Circular buffer for backpressure, rolling file appender, distributed log aggregation (ELK/etc.)",
        "signal": "Async logging: caller enqueues fast, worker thread drains and fans out. LogLevel filter gates BEFORE queue insertion — most logs never enter the queue. CopyOnWriteArrayList for appender list = lock-free fan-out (read-heavy access pattern).",
        "traps": [
            "Synchronous I/O in caller's hot path (blocks request thread)",
            "Unbounded queue → OOM under burst",
            "ArrayList of appenders → ConcurrentModificationException on hot-add",
        ],
        "soundbite": "log() enqueues to BlockingQueue, worker thread drains and fans out to appenders via Strategy. Filters Decorate appenders. CopyOnWriteArrayList for appender registry — lock-free reads. Backpressure policy on queue full: drop-newest vs. block-caller — Strategy.",
    },
    "ratelimiter": {
        "entities": "RateLimiter (interface), TokenBucket, FixedWindow, SlidingWindow (Strategies), KeyedRateLimiter (per-user via ConcurrentHashMap)",
        "patterns_day1": "Strategy (algorithm choice), per-key ConcurrentHashMap.computeIfAbsent for bucket creation",
        "patterns_step5": "Distributed rate limiting via Redis Lua script; sliding-window log (precise) vs counter (approximate) tradeoff",
        "signal": "Per-key state via ConcurrentHashMap.computeIfAbsent — atomic creation, no lost-update race. Token bucket uses System.nanoTime (monotonic) — refill happens LAZILY on check, not via background thread.",
        "traps": [
            "Single shared bucket — kills per-user fairness",
            "Time arithmetic in doubles or wall-clock millis (wall clock can go backwards on NTP sync)",
            "Background refill thread — pure overhead vs. lazy refill",
        ],
        "soundbite": "Strategy chooses algorithm (token-bucket / fixed-window / sliding-window). computeIfAbsent creates per-key buckets atomically. nanoTime for monotonic time. Lazy refill — calculate tokens-to-add on every check, no scheduler needed. Distributed version → Redis Lua script.",
    },
    "inventory": {
        "entities": "Product, Warehouse, StockEntry (productId+warehouseId → quantity), Reservation (HELD/CONFIRMED/RELEASED/EXPIRED), InventoryService",
        "patterns_day1": "Per-warehouse synchronized, state machine for Reservation, Observer for low-stock alerts",
        "patterns_step5": "Cross-warehouse transfers (saga), multi-tenant isolation, eventual consistency under partition",
        "signal": "Reservation-with-TTL: HELD on add-to-cart → CONFIRMED on order-placed. Per-warehouse synchronized on StockEntry mutation. Stale HELDs auto-expire via ScheduledExecutor — released back to available.",
        "traps": [
            "Global lock across warehouses (cross-warehouse interference)",
            "Quantity goes negative under race (no atomic check-then-decrement)",
            "Not expiring stale reservations — phantom stock locks",
        ],
        "soundbite": "Reservations HELD with TTL then CONFIRMED on checkout. Per-warehouse synchronized — cross-warehouse traffic doesn't contend. Observer fires when stock crosses low-threshold. Scheduled job expires stale holds. Cross-warehouse transfer is a saga — out of scope for v1.",
    },
    "notification": {
        "entities": "Notification, Recipient, Channel (Email/SMS/Push — Strategy), NotificationService, Subscription, Topic",
        "patterns_day1": "Strategy (Channel), Observer (Topic→Subscribers), Factory (channel registry), async fan-out",
        "patterns_step5": "Retry with exponential backoff, templating (Decorator), deduplication, priority queues for SMS vs marketing",
        "signal": "CopyOnWriteArrayList for topic subscribers — read-heavy fan-out, lock-free. Per-channel async send + retry-with-backoff via ScheduledExecutorService. Channel chosen by recipient preference + Strategy.",
        "traps": [
            "Synchronous send blocks publisher thread (use async)",
            "ArrayList of subscribers — concurrent mod exception",
            "No retry on transient channel failure — silent drops",
        ],
        "soundbite": "Topic has CopyOnWriteArrayList of subscribers — lock-free fan-out. Each subscriber's preferred Channel chosen by Strategy. Async send + retry with exponential backoff via ScheduledExecutorService. Templating is a Decorator on the message body.",
    },
    "jobscheduler": {
        "entities": "Job (priority, runAt, payload, status), JobScheduler, Worker pool, JobStatus state machine, PriorityBlockingQueue",
        "patterns_day1": "State machine (enum), Strategy for retry policy, Producer-Consumer (queue + workers)",
        "patterns_step5": "Persistence (job survives restart), cron expressions, circuit breaker per job-type, distributed via Kafka/SQS",
        "signal": "PriorityBlockingQueue orders by (priority DESC, runAt ASC). Workers drain. Cancellation flips state — running job checks isInterrupted between phases. Delayed jobs via DelayQueue OR timestamp-gate in priority comparator.",
        "traps": [
            "Worker death leaves job stuck RUNNING — need heartbeat / lease",
            "Priority inversion — low-priority hogs worker, high-priority starves",
            "Unbounded queue → OOM",
        ],
        "soundbite": "PriorityBlockingQueue orders jobs by priority + runAt. Worker pool drains. State machine in enum guards transitions. Retry Strategy — fixed delay or exponential. Cancellation flips state; running job checks isInterrupted between steps and exits cleanly.",
    },
    "meetingscheduler": {
        "entities": "Meeting (start, end, attendees, room), User, Room, Calendar (TreeMap<start, Meeting>), MeetingService",
        "patterns_day1": "TreeMap.floor/ceiling for O(log n) overlap detection, half-open intervals [start, end), Clock injection",
        "patterns_step5": "Recurring events (RRULE), time-zone handling (ZonedDateTime), conflict-resolution Strategy",
        "signal": "Overlap detection in O(log n) — TreeMap.floor(newStart) gives the only meeting that can overlap. If its end > newStart, conflict. Half-open intervals [start, end) eliminate boundary bugs (end=10am AND start=10am don't overlap).",
        "traps": [
            "Linear scan O(n) for conflict detection",
            "Closed intervals [start, end] — back-to-back meetings flagged as overlap",
            "Wall-clock new Date() — untestable; inject Clock",
        ],
        "soundbite": "Calendar is a TreeMap of meetings keyed by start. floor(newStart) gives the only adjacent candidate; if its end > newStart, conflict. Half-open intervals [start, end). Per-room synchronized for the book CAS. Clock injected — tests use Clock.fixed(...).",
    },
    "paymentgateway": {
        "entities": "Payment (PENDING/AUTHORIZED/CAPTURED/REFUNDED/FAILED), PaymentProcessor (per-gateway Strategy), IdempotencyKey, PaymentRequest",
        "patterns_day1": "State machine (enum + EnumMap<Status, EnumSet<Status>>), Strategy (per-gateway), idempotency via computeIfAbsent",
        "patterns_step5": "Webhook callback handling, 3D-Secure flow, reconciliation, chargebacks, ledger",
        "signal": "Idempotency: key → ConcurrentHashMap.computeIfAbsent on payment registry. Duplicate call with same key returns the EXISTING payment object (no double charge). State machine forbids skipping — PENDING → CAPTURED illegal without AUTHORIZE.",
        "traps": [
            "No idempotency → double charge on retry",
            "doubles for money",
            "Allowing capture without prior authorization",
        ],
        "soundbite": "Idempotency key → computeIfAbsent on payment registry returns existing payment on duplicate. State machine in enum forbids invalid transitions (PENDING→CAPTURED skipped AUTHORIZED). Per-gateway PaymentProcessor Strategy. Money in long cents always.",
    },
    "urlshortener": {
        "entities": "UrlShortener (facade), Base62Encoder, IdGenerationStrategy (Counter/Random), urlToCode + codeToUrl maps",
        "patterns_day1": "Strategy (ID generation), Facade, Base62 encoding (URL-safe), idempotency via urlToCode.computeIfAbsent",
        "patterns_step5": "Custom alias, expiration TTL, click analytics, distributed counter via Snowflake or DB sequence",
        "signal": "Bidirectional map (urlToCode + codeToUrl). Idempotency via urlToCode.computeIfAbsent — same long URL always returns the same code. Counter strategy uses AtomicLong + Base62; Random strategy retries on collision via Predicate<String> isAvailable.",
        "traps": [
            "Forward map only — same URL shortened twice produces two codes",
            "UUID for code — not URL-safe-pretty and too long",
            "Random strategy without bounded retry — infinite loop on near-full namespace",
        ],
        "soundbite": "Idempotent shorten via urlToCode.computeIfAbsent — same URL → same code. AtomicLong + Base62 for Counter strategy (monotonic, predictable). Random strategy retries on collision with a bounded loop. Bidirectional maps for O(1) resolve in both directions.",
    },
    "filesystem": {
        "entities": "Node (abstract), File extends Node, Directory extends Node (children list), Path, FileSystemService",
        "patterns_day1": "Composite (Directory contains Nodes), polymorphism on Node.size() / Node.search(), Tell-Don't-Ask",
        "patterns_step5": "Find command (Filter Strategy + composite AND/OR/NOT), permissions, symlinks, watchers (Observer)",
        "signal": "Directory.size() recurses through children polymorphically — Files return file size, Directories sum children. NO caller-side branching on isFile/isDirectory. That's the entire payoff of the Composite pattern.",
        "traps": [
            "Caller branches on isFile vs isDirectory instead of polymorphic call",
            "Copy-paste size aggregation logic at each call site",
            "Not handling cycles if symlinks added (need visited-set)",
        ],
        "soundbite": "Composite — Directory contains Nodes (Files OR Directories). size() recurses polymorphically. Visitor for traversal. Linux find = FilterStrategy composition layered on top — AND/OR/NOT for name/size/extension filters. No type-switching in callers.",
    },
    "cabbooking": {
        "entities": "Rider, Driver (status: AVAILABLE/ON_TRIP), Ride (state machine: REQUESTED→MATCHED→IN_PROGRESS→COMPLETED), Location, DriverMatchingStrategy, PricingStrategy",
        "patterns_day1": "Strategy (matching + pricing), enum state machine (Ride lifecycle), optimistic-match loop (ranked list + per-driver synchronized CAS)",
        "patterns_step5": "Spatial index (quadtree / H3), driver-can-reject (two-phase commit), surge from external demand monitor, multiple cab types",
        "signal": "MatchingStrategy returns a RANKED LIST, not a single driver. Service iterates calling Driver.tryReserve() (synchronized atomic AVAILABLE→ON_TRIP). First to win the CAS takes the ride. Zero double-booking under contention — empirically proven in driver scenario 6 (50 riders × 10 drivers).",
        "traps": [
            "Matching returns single driver → no fallback if CAS lost",
            "Holding a lock during strategy.rankCandidates (serializes matching)",
            "Cancel-after-match path doesn't release driver — phantom ON_TRIP forever",
        ],
        "soundbite": "MatchingStrategy ranks; service walks the ranked list calling Driver.tryReserve — synchronized atomic AVAILABLE → ON_TRIP. Loser of the race tries the next candidate. PricingStrategy with surge in basis points (passed at call time). Enum state machine for Ride. Optimistic match scales without global locks.",
    },
    "chess": {
        "entities": "Board, Piece (abstract, with Pawn/Knight/Bishop/Rook/Queen/King subclasses), Color, Position (record), Move (record), ChessGame",
        "patterns_day1": "Polymorphism on Piece.isValidMove (Board NEVER switches on piece type), 3-layer legality (Piece→Board→Game)",
        "patterns_step5": "Castling / en-passant / promotion (special-case methods on King + Pawn), FEN parsing, perft testing, PGN export",
        "signal": "King-safety = try-move + isSquareAttackedBy + undo. Board has applyMove + undoMove primitives. ChessGame uses them as a 'peek' to verify your own king isn't attacked after the move. Checkmate = inCheck && noLegalMove; Stalemate = !inCheck && noLegalMove.",
        "traps": [
            "switch (piece.type) in Board — kills polymorphism, hard to extend",
            "Not undoing the peek-move (corrupts board state)",
            "Forgetting pinned-piece detection (move that exposes own king is illegal)",
        ],
        "soundbite": "Piece.isValidMove is polymorphic — Board never branches on piece type. Three legality layers: Piece (geometry) → Board (physical state, path clear) → ChessGame (turn order + king safety + checkmate). King safety = try-move + peek-attacked + undo. Checkmate = no legal move AND in check.",
    },
    "elevator": {
        "entities": "Elevator (state: IDLE/MOVING_UP/MOVING_DOWN), Request (floor + direction), ElevatorSystem, DispatchStrategy (Nearest / LOOK), Direction",
        "patterns_day1": "State machine, Strategy (DispatchStrategy — pre-baked because multiple algorithms are guaranteed), per-elevator request queue",
        "patterns_step5": "Multi-elevator coordination, peak-hour vs off-peak strategy switching, priority floors (lobby), emergency override",
        "signal": "DispatchStrategy assigns hall calls to elevators (which one?); per-elevator state machine drives motion (what next?). LOOK-style scan — serve all requests in current direction before reversing. Hall vs. car buttons go through different paths.",
        "traps": [
            "Single shared queue for all elevators (defeats parallelism)",
            "Elevator reverses mid-direction (real elevators don't)",
            "Treating hall calls and car calls identically",
        ],
        "soundbite": "DispatchStrategy assigns hall calls to elevators. Per-elevator state machine (IDLE → MOVING_UP → IDLE). LOOK-scan — serve everything in the current direction first, then reverse. Hall calls go through dispatch; car calls go straight to the chosen elevator's queue. Multiple strategies — Nearest, SCAN, LOOK — chosen by time-of-day.",
    },
}


# ----------------------- Reference pages content -----------------------
REFERENCE_PAGES = [
    ("Pattern → Problem Index", """
## Strategy (multiple implementations chosen at runtime)
- **Snake & Ladder** — Dice (real + test-fake)
- **Logging Service** — LogAppender (console/file/remote), LogFilter
- **Rate Limiter** — algorithm choice (TokenBucket / FixedWindow / SlidingWindow)
- **Notification System** — Channel (Email / SMS / Push)
- **Splitwise** — SplitStrategy (Equal / Exact / Percentage)
- **Job Scheduler** — retry policy
- **Payment Gateway** — per-gateway PaymentProcessor
- **URL Shortener** — IdGenerationStrategy (Counter / Random)
- **Cab Booking** — DriverMatchingStrategy + PricingStrategy (the big one)
- **Elevator** — DispatchStrategy (Nearest / LOOK / SCAN)
- **Movie Ticket Booking** — PricingStrategy
- **Amazon Locker** — size-matching strategy

## State Pattern — class-per-state (GoF)
- **Vending Machine** — NoCoinState, HasCoinState, DispensingState (the textbook example)

## State Machine — enum-driven (transition map)
- **Payment Gateway** — PENDING / AUTHORIZED / CAPTURED / REFUNDED / FAILED
- **Job Scheduler** — QUEUED / RUNNING / COMPLETED / FAILED / CANCELLED
- **Inventory** — Reservation: HELD / CONFIRMED / RELEASED / EXPIRED
- **Cab Booking** — Ride: REQUESTED / MATCHED / IN_PROGRESS / COMPLETED / CANCELLED
- **Movie Ticket** — Seat: AVAILABLE / HELD / BOOKED

## Observer / Pub-Sub
- **Notification System** — Topic → Subscribers (canonical example)
- **Inventory** — low-stock threshold alerts
- **Movie Ticket** — booking confirmation events
- **Logging Service** — multi-appender fan-out (Observer-ish)

## Factory
- **Notification System** — channel factory by name
- **Logging Service** — appender factory from config

## Composite
- **File System** — Directory contains Nodes (Files or Directories)

## Decorator
- **Logging Service** — Filter wraps Appender
- **Notification System** — templating wraps message body

## Facade
- Almost every *Service class. Most explicit examples: **URL Shortener**, **Payment Gateway**, **Cab Booking Service**, **Inventory Service**

## Polymorphism (no pattern — just OOP done right)
- **Chess** — Piece.isValidMove dispatched by type (Pawn/Knight/.../King), Board NEVER branches on type
- **File System** — Node.size() / Node.search() dispatched by File vs. Directory

## Patterns I did NOT pre-bake in any problem
- **Singleton** — instantiation choice, not design. Vending Machine's 3 state objects are *de facto* singletons but never declared as such.
- **Builder** — mentioned for Meeting / complex constructions but never required.
- **Prototype** — no use case.
- **Chain of Responsibility** — could fit Notification retry escalation; deferred to Step 5.
"""),

    ("Concurrency Idiom → Problem Index", """
## ConcurrentHashMap.computeIfAbsent — atomic per-key creation
- **URL Shortener** — `urlToCode.computeIfAbsent` makes shorten() idempotent
- **Rate Limiter** — per-key bucket created exactly once even under burst
- **Inventory** — per-product StockEntry creation
- **Payment Gateway** — idempotency-key lookup (duplicate request returns existing payment)

## Per-resource synchronized CAS (atomic state transition on one entity)
- **Parking Lot** — Spot.tryReserve (AVAILABLE → OCCUPIED)
- **Amazon Locker** — Locker.tryReserve
- **Movie Ticket Booking** — Seat hold via synchronized seat
- **Cab Booking** — Driver.tryReserve (AVAILABLE → ON_TRIP)
- **Inventory** — per-warehouse synchronized on StockEntry mutation

## CopyOnWriteArrayList — lock-free fan-out for read-heavy lists
- **Notification System** — topic subscribers
- **Logging Service** — appender registry

## PriorityBlockingQueue — concurrent priority queue
- **Job Scheduler** — jobs ordered by (priority DESC, runAt ASC)

## ScheduledExecutorService — delayed / periodic tasks
- **Job Scheduler** — delayed jobs
- **Notification System** — retry with exponential backoff
- **Movie Ticket Booking** — hold timeout / auto-release
- **Inventory** — reservation expiry sweeper

## DelayQueue — alternative to ScheduledExecutorService for timed dequeue
- **Job Scheduler** — delayed-job variant

## Ordered lock acquisition (deadlock prevention)
- **Splitwise** — sort the two user IDs before acquiring locks on cross-payment

## AtomicLong / AtomicInteger — lock-free counters
- **URL Shortener** — monotonic ID generation
- **Job Scheduler** — unique job IDs

## BlockingQueue + worker thread — async producer-consumer
- **Logging Service** — log calls enqueue, worker drains
- **Notification System** — async send
- **Job Scheduler** — workers drain jobs

## CountDownLatch — used in DRIVER TESTS to release a burst of threads simultaneously
- **Cab Booking**, **Movie Ticket**, **Inventory**, **URL Shortener** — concurrency proof scenarios

## Monotonic time — System.nanoTime() (not currentTimeMillis())
- **Rate Limiter** — token-bucket refill (wall clock can move backwards on NTP sync)
- **Job Scheduler** — duration measurement
"""),

    ("Company → Priority Problems", """
> Use this to prioritize when you have a specific interview lined up.

## Adobe (Noida) — SDE-2
Adobe loves classic LLDs with clean OOP and Strategy/State patterns.
1. **Parking Lot** — universal warm-up; tests basics
2. **Snake & Ladder** — frequent simple problem
3. **Splitwise** — money + graph thinking
4. **Movie Ticket Booking** — BookMyShow flavor
5. **File System** — Linux `find` command variant
6. **Vending Machine** — State pattern showcase
7. **Inventory / Library Management** — entity modeling

## Microsoft India — SDE-2
Microsoft leans HLD-heavy but LLD comes up. Focus: producer-consumer, caching, async.
1. **Vending Machine** — State pattern
2. **Logging Service** — async, appenders, filters
3. **LRU Cache** — data structures
4. **Job Scheduler** — Producer-Consumer pattern
5. **Notification System** — Pub-Sub
6. **Rate Limiter** — Strategy + concurrency

## Amazon India — SDE-2
Amazon loves business-domain LLDs with concurrency correctness.
1. **Parking Lot** — must-do
2. **Splitwise** — money atomicity
3. **Movie Ticket Booking** — BookMyShow / seat hold
4. **Amazon Locker** — internal-product flavor
5. **Inventory Management** — warehouse / fulfillment context
6. **Cab Booking** — multi-pattern showcase
7. **Rate Limiter** — API design context

## Universal warm-ups (any company, in any order)
- Parking Lot, Vending Machine, Snake & Ladder, LRU Cache, Connect Four

## Multi-pattern flagships (study last — these are 45-min-tight)
- Cab Booking, Chess, Elevator, File System, URL Shortener
"""),

    ("45-Minute Round Playbook", """
> Generic schedule that applies to every problem in this deck.

## Time Budget (Hello Interview's framework)
| Step | Time | What you produce |
|------|------|------------------|
| 1. Requirements clarification | 0 – 5 min | Functional list, non-functional list, explicit out-of-scope list |
| 2. Entity identification | 5 – 10 min | Entity table with purpose + lifetime + key invariants |
| 3. Class design | 10 – 20 min | Key classes + signatures + brief explanation of design seams |
| 4. Core implementation | 20 – 35 min | Code the critical happy path (NOT all CRUD) |
| 5. Extensibility / deep-dive | 35 – 45 min | Answer "how would you extend X?" — show senior signal |

## What to Say at Minute 5
- "Let me restate the requirements I'm building against..."
- "What's explicitly out of scope: persistence, payments, network, UI..."
- "Non-functional: concurrent requests, money discipline, in-memory, deterministic-when-clock-injected"
- "Out-of-scope upfront saves us 10 minutes of confusion later"

## What to Say at Minute 15
- "The key invariant is [no double-booking / no negative stock / atomic money transfer]"
- "I'll pre-bake [Strategy / State machine] because the one-sentence test passes — I need at least 2 implementations on day 1"
- "I'm deferring [Observer / Decorator / Factory] until Step 5 because there's only one consumer right now"

## What to Say at Minute 30 (during implementation)
- "Let me code the critical happy path first — [describe]"
- "The contention point is [identify race]; I'm using [synchronized / computeIfAbsent / ranked-list optimistic match] to make it atomic"
- "I'll show the test scenario that proves no double-booking"

## What to Say at Minute 40 (extensibility deep-dive)
- "Three escalating answers: simple in-memory → with thread-safety → distributed"
- "The Strategy seam absorbs this without any callsite change"
- "I'd add Observer here if the requirement evolves to [...]"

## Senior Signals to Explicitly Verbalize
- "Tell, don't ask — the object decides, callers don't peek-and-branch"
- "Money is `long cents`. Percentages are basis points (10000 = 100%)"
- "Half-open intervals `[start, end)` to avoid boundary bugs"
- "Clock is injected for testability — tests use `Clock.fixed(...)`"
- "One-sentence test for Strategy: do I need 2+ implementations on day 1? If yes, pre-bake; if no, defer"
- "Class-per-state when behavior differs per state; enum-state-machine when transitions are bookkeeping"
- "Per-resource lock granularity, never a global lock"
- "Atomic check-and-set via synchronized or computeIfAbsent — never read-then-write without atomicity"
- "Empirical concurrency proof: 50 threads, CountDownLatch barrier, assert invariant holds"

## Red Flags (avoid saying these)
- "I'll make it a Singleton" — that's an instantiation choice, not a pattern
- "I'll use a HashMap for synchronization" — use ConcurrentHashMap or explicit locks
- "I'll just use doubles for money for now" — never; long cents always
- "I'll add Observer/Factory/Builder in case we need it" — YAGNI; introduce on Step 5 demand

## When You're Stuck — Recovery Tactics
- **Drawing blank on entities** — describe the domain in plain English; nouns become entities
- **Stuck on a race condition** — name the invariant ("no two riders match the same driver"); the lock falls out of the invariant
- **Pattern doesn't fit** — say so; "Strategy would over-engineer this — an if-else is the right granularity for now"
- **Out of time** — finish the happy path; explicitly defer extensions: "given more time I'd add..."
"""),

    ("Anti-Pattern Catalog (deduped from all 20 problems)", """
> Read this the night before — these are the mistakes that show up across the deck.

## Money / Currency
- **Using `double` or `float` for money** — silent rounding, accumulates errors
- **Symmetric rounding** in equal-split remainders — sum no longer matches total
- **Percentages as `double` instead of basis points (int)** — 10% as 0.1 vs. 1000 bps

## Concurrency
- **Global lock on a registry** — kills throughput; use per-resource locks
- **Per-resource locks acquired in arbitrary order** — deadlock on cross-resource ops (Splitwise lesson)
- **Read-then-write without atomicity (TOCTOU)** — two callers read the same value, both write; use `computeIfAbsent` or synchronized CAS
- **`ArrayList` of subscribers/observers** — `ConcurrentModificationException` on hot add/remove; use `CopyOnWriteArrayList`
- **Unbounded queues** — OOM under burst; bounded with explicit backpressure policy
- **Synchronous I/O in hot path** — logger / notification publishers must be async
- **Wall clock for time math** — `currentTimeMillis()` can go backwards on NTP sync; use `nanoTime()` for monotonic

## State Machines
- **Direct setter on status field** — bypasses transition validation; force mutation through named lifecycle methods (match/start/complete/cancel)
- **Skipping intermediate states** — REQUESTED → COMPLETED without MATCHED/IN_PROGRESS
- **No allowed-transitions validation** — invalid transitions silently succeed
- **enum-state-machine where class-per-state belongs** — when behavior differs per state, not just data
- **class-per-state where enum belongs** — when transitions are pure bookkeeping with no behavior difference

## Design Seams
- **switch on type instead of polymorphism** — especially Chess Board switching on Piece type
- **Caller-side branching on enum/type** — push the decision into the object (Tell, Don't Ask)
- **Strategy pattern with one implementation** — over-engineering; add second implementation OR collapse to a method
- **Single-method classes that should just be functions** — over-OOP

## Time / Spatial
- **Closed intervals `[start, end]`** — back-to-back meetings false-flag as overlap
- **Wall-clock `new Date()` / `Instant.now()` without `Clock` injection** — tests can't deterministically reproduce
- **Linear scan over all entities for nearest/conflict** — use spatial index (Cab Booking) or `TreeMap.floor` (Meeting Scheduler)

## Testing
- **Random behavior without seed/fake-injection** — flaky tests (use FixedSequenceDice, fixed Clock)
- **No empirical concurrency proof** — just "looks thread-safe" is not enough; show 50-thread test with CountDownLatch barrier asserting the invariant
- **Asserting on private state via reflection** — assert on public behavior; if you can't, the API is wrong

## Out-of-Scope Discipline
- **Pre-baking patterns "in case we need them"** — YAGNI; one-sentence test for each pattern
- **Adding fields you might use** — start with what the requirements demand
- **Designing for distributed scale on day 1** — scope is in-memory; mention distributed in Step 5

## Communication
- **Saying "I'll make it a Singleton"** — instantiation choice, not pattern
- **Naming patterns without showing them solve a real problem** — "I'd use Decorator" without explaining what gets wrapped why
- **Ignoring the interviewer's hints** — if they say "what about concurrency?" — pivot, don't continue with CRUD
"""),
]


# ============================================================================
#                            DOCX RENDERING HELPERS
# ============================================================================

# ---------------- Word TOC field (auto-paginates on open) ----------------

def add_toc_field(paragraph):
    """Insert a Word TOC field. User opens in Word, hits F9 to populate."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "  Right-click and choose 'Update Field' (or press F9) to populate page numbers."
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    r_element = run._r
    for el in (fldChar1, instrText, fldChar2, placeholder, fldChar3):
        r_element.append(el)


# ---------------- Inline formatting ----------------

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"   # bold
    r"|(`[^`]+`)"        # inline code
    r"|(\*[^*]+\*)"      # italic
)

def add_inline_runs(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(10)
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------- Code block (gray-shaded table cell) ----------------

def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), "808080")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear"); shading.set(qn("w:color"), "auto"); shading.set(qn("w:fill"), "F2F2F2")
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9)


# ---------------- Tables ----------------

def parse_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(c) <= set("-:") for c in rows[1]):
        rows.pop(1)
    return rows

def add_md_table(doc, rows):
    if not rows: return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


# ---------------- MD parser ----------------

def render_markdown_to_doc(doc, md_text, heading_offset=0):
    """heading_offset: shift all heading levels by N (so MD H1 becomes Word H(1+offset))."""
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf); code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        if re.match(r"^---+\s*$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "999999")
            pBdr.append(bot); pPr.append(pBdr)
            i += 1; continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1)) + heading_offset
            text = m.group(2).strip()
            p = doc.add_heading(text, level=min(level, 5))
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1; continue

        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i+1]):
            tbl_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                tbl_lines.append(lines[i]); i += 1
            rows = parse_table(tbl_lines)
            add_md_table(doc, rows); continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)); text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.4)
            add_inline_runs(p, text)
            i += 1; continue

        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1)); text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.4)
            add_inline_runs(p, text)
            i += 1; continue

        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(text); r.italic = True
            i += 1; continue

        if not line.strip():
            i += 1; continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1


# ---------------- Quick Card ----------------

def add_quick_card(doc, problem_title, qc):
    """One-page distilled card before the full walkthrough."""
    # Card title
    h = doc.add_heading("Quick Card", level=2)
    for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)

    # Wrap card content in a 1-cell shaded table for visual separation
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    # subtle border + light shade
    tbl = table._tbl; tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:color"), "666666")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear"); shading.set(qn("w:color"), "auto"); shading.set(qn("w:fill"), "FAFAFA")
    cell._tc.get_or_add_tcPr().append(shading)

    # clear default empty paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    def label_row(label, body, mono=False):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label + "  ")
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        add_inline_runs(p, body)
        for run in p.runs[1:]:
            run.font.size = Pt(10)
            if mono: run.font.name = "Consolas"

    label_row("ENTITIES",         qc["entities"])
    label_row("PATTERNS (Day 1)", qc["patterns_day1"])
    label_row("PATTERNS (Step 5)",qc["patterns_step5"])
    label_row("SENIOR SIGNAL",    qc["signal"])

    # Top traps as bullets
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TOP TRAPS"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for trap in qc["traps"]:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.add_run("•  ").font.size = Pt(10)
        add_inline_runs(p, trap)
        for run in p.runs[1:]:
            run.font.size = Pt(10)

    # Soundbite — italic quote box
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SOUNDBITE"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_before = Pt(0)
    r = p.add_run('"' + qc["soundbite"] + '"')
    r.italic = True; r.font.size = Pt(10)


# ============================================================================
#                                MAIN BUILD
# ============================================================================

def build():
    doc = Document()

    # ---- Document defaults ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)
    for hn in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"):
        s = doc.styles[hn]; s.font.color.rgb = RGBColor(0, 0, 0); s.font.name = "Calibri"
    for section in doc.sections:
        section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

    # ---- Title page ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Low-Level Design"); r.bold = True; r.font.size = Pt(28)
    for _ in range(2): doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Interview Revision Sheet"); r.font.size = Pt(20)
    for _ in range(2): doc.add_paragraph()
    sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(f"{len(STUDY_SEQUENCE)} Problems  •  Adobe / Microsoft / Amazon India SDE-2")
    r.italic = True; r.font.size = Pt(12)
    for _ in range(4): doc.add_paragraph()
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Study sequence: Foundation → Core → Advanced → Multi-pattern Flagships")
    r.font.size = Pt(11); r.italic = True
    for _ in range(2): doc.add_paragraph()
    note2 = doc.add_paragraph(); note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note2.add_run("Each problem: Quick Card (1 page) + Full Walkthrough · Reference pages at the back")
    r.font.size = Pt(10); r.italic = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- Index / TOC ----
    h = doc.add_heading("Contents", level=1)
    for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)

    if STATIC_INDEX_PAGES:
        # Static pre-paginated index — no F9 needed
        # Build two parallel rows: problem-index (first 20) and reference-index (last 5)
        problem_entries   = [(k, v) for k, v in STATIC_INDEX_PAGES.items() if re.match(r"^\d+\.", k)]
        reference_entries = [(k, v) for k, v in STATIC_INDEX_PAGES.items() if not re.match(r"^\d+\.", k)]

        # Problem table
        intro = doc.add_paragraph()
        r = intro.add_run("Twenty problems in study order. Each one: 1-page Quick Card + Full Walkthrough.")
        r.italic = True; r.font.size = Pt(10)

        tbl = doc.add_table(rows=len(problem_entries) + 1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        # header
        hdr = tbl.rows[0]
        for ci, label in enumerate(["Problem", "Page"]):
            cell = hdr.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            run = p.add_run(label); run.bold = True; run.font.size = Pt(10)
        # rows
        for i, (name, page) in enumerate(problem_entries, start=1):
            row = tbl.rows[i]
            for ci, val in enumerate([name, str(page)]):
                cell = row.cells[ci]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                run = p.add_run(val); run.font.size = Pt(10)
                if ci == 1:  # page-number column right-aligned
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Reference table — separate, below
        doc.add_paragraph()
        intro2 = doc.add_paragraph()
        r = intro2.add_run("Cross-cutting reference pages — pattern index, concurrency idiom map, company priorities, 45-min playbook, anti-pattern catalog.")
        r.italic = True; r.font.size = Pt(10)

        tbl2 = doc.add_table(rows=len(reference_entries) + 1, cols=2)
        tbl2.style = "Table Grid"
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = tbl2.rows[0]
        for ci, label in enumerate(["Reference Page", "Page"]):
            cell = hdr.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            run = p.add_run(label); run.bold = True; run.font.size = Pt(10)
        for i, (name, page) in enumerate(reference_entries, start=1):
            row = tbl2.rows[i]
            for ci, val in enumerate([name, str(page)]):
                cell = row.cells[ci]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                run = p.add_run(val); run.font.size = Pt(10)
                if ci == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        # Fallback: Word TOC field placeholder (needs F9 in Word)
        note = doc.add_paragraph()
        r = note.add_run("Open this file in Word. To populate page numbers in the table below, click anywhere inside the placeholder and press F9 (or right-click → 'Update Field').")
        r.italic = True; r.font.size = Pt(10)
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_toc_field(p)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- Each problem: Heading + Tier + Quick Card + Walkthrough ----
    current_tier = None
    for idx, (slug, title_text, tier) in enumerate(STUDY_SEQUENCE, start=1):
        md_path = ROOT / slug / "INTERVIEW_WALKTHROUGH.md"
        if not md_path.exists():
            print(f"  WARN: skipping {slug}"); continue

        # Problem title (Heading 1 → TOC entry)
        h = doc.add_heading(f"{idx}. {title_text}", level=1)
        for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)

        # Tier badge under the title
        p = doc.add_paragraph()
        r = p.add_run(tier); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Quick Card
        qc = QUICK_CARDS.get(slug)
        if qc:
            add_quick_card(doc, title_text, qc)
        else:
            print(f"  WARN: no quick-card for {slug}")

        # Page break before full walkthrough
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Full walkthrough as a clearly labeled sub-section
        h = doc.add_heading("Full Walkthrough", level=2)
        for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)

        md_text = md_path.read_text(encoding="utf-8")
        # drop the MD's own H1 (problem title — we already showed it)
        md_text = re.sub(r"^#\s+[^\n]+\n+", "", md_text, count=1)
        # MD subsections (##, ###) are demoted by 1 to sit under our H2 'Full Walkthrough'
        render_markdown_to_doc(doc, md_text, heading_offset=1)

        # Page break between problems
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        print(f"  added: {idx:2d}. {title_text}")

    # ---- Reference pages ----
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sep.add_run("— REFERENCE PAGES —"); r.bold = True; r.font.size = Pt(14)
    for _ in range(2): doc.add_paragraph()
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Cross-cutting indexes to read the night before — pattern lookup, concurrency idiom map, company priorities, 45-min playbook, anti-patterns.")
    r.italic = True; r.font.size = Pt(11)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for ref_title, ref_md in REFERENCE_PAGES:
        h = doc.add_heading(ref_title, level=1)
        for run in h.runs: run.font.color.rgb = RGBColor(0, 0, 0)
        render_markdown_to_doc(doc, ref_md.strip(), heading_offset=1)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        print(f"  added reference: {ref_title}")

    doc.save(OUTPUT)
    print(f"\nWrote {OUTPUT}")
    print(f"({OUTPUT.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    build()
